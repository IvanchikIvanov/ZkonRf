from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bot.services.document_templates.models import DraftState
from bot.services.document_templates.templates import get_template


HEADINGS = {
    "seller_refund_claim": "ПРЕТЕНЗИЯ\nо возврате денежных средств за товар",
    "poor_service_claim": "ПРЕТЕНЗИЯ\nпо договору оказания услуг",
    "rospotrebnadzor_complaint": "ЖАЛОБА\nо нарушении прав потребителя",
    "consumer_lawsuit": "ИСКОВОЕ ЗАЯВЛЕНИЕ\nо защите прав потребителя",
    "bank_mfo_collector_claim": "ЗАЯВЛЕНИЕ\nо защите прав заемщика/потребителя финансовых услуг",
    "transport_complaint": "ЖАЛОБА\nна нарушение прав пассажира",
    "education_refund_claim": "ЗАЯВЛЕНИЕ\nо возврате денежных средств за обучение",
}


def _value(draft: DraftState, field_id: str, default: str = "") -> str:
    return draft.fields.get(field_id, default).strip()


def _add_paragraph(document: Document, text: str, bold: bool = False, center: bool = False) -> None:
    paragraph = document.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold


def render_document(draft: DraftState, output_dir: Path) -> Path:
    template = get_template(draft.template_id)
    output_dir.mkdir(parents=True, exist_ok=True)

    document = Document()
    _add_paragraph(document, f"Кому: {_value(draft, 'recipient_name')}", bold=True)
    if _value(draft, "recipient_address"):
        _add_paragraph(document, f"Адрес: {_value(draft, 'recipient_address')}")
    _add_paragraph(document, f"От: {_value(draft, 'user_full_name')}", bold=True)
    _add_paragraph(document, f"Адрес: {_value(draft, 'user_address')}")
    _add_paragraph(document, f"Контакт: {_value(draft, 'user_contact')}")
    document.add_paragraph()

    _add_paragraph(document, HEADINGS.get(draft.template_id, template.title).upper(), bold=True, center=True)
    document.add_paragraph()

    _add_paragraph(document, f"Дата события: {_value(draft, 'event_date')}")
    if _value(draft, "amount"):
        _add_paragraph(document, f"Сумма: {_value(draft, 'amount')}")

    if draft.template_id == "consumer_lawsuit":
        _add_paragraph(document, f"Суд: {_value(draft, 'court_name')}")
        _add_paragraph(document, f"Ответчик: {_value(draft, 'defendant_name')}")
        _add_paragraph(document, f"Цена иска: {_value(draft, 'claim_amount')}")
        _add_paragraph(document, f"Досудебная претензия: {_value(draft, 'pretrial_claim')}")

    if draft.template_id == "bank_mfo_collector_claim":
        _add_paragraph(document, f"Номер договора/займа: {_value(draft, 'contract_number')}")
        _add_paragraph(document, f"Спорная сумма: {_value(draft, 'disputed_amount')}")

    document.add_paragraph()
    _add_paragraph(document, "Обстоятельства", bold=True)
    _add_paragraph(document, _value(draft, "facts"))

    document.add_paragraph()
    _add_paragraph(document, "Требования", bold=True)
    _add_paragraph(document, _value(draft, "demand"))

    if _value(draft, "evidence"):
        document.add_paragraph()
        _add_paragraph(document, "Доказательства", bold=True)
        _add_paragraph(document, _value(draft, "evidence"))

    document.add_paragraph()
    _add_paragraph(document, "Приложения: копии подтверждающих документов при наличии.")
    _add_paragraph(document, "Дата: ____________________")
    _add_paragraph(document, "Подпись: ____________________")

    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"{template.filename_prefix}_{draft.user_id}_{timestamp}.docx"
    path = output_dir / filename
    document.save(str(path))
    return path
